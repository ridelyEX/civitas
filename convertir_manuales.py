#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para convertir los manuales de usuario de Markdown a Word y PDF
Autor: Sistema de Documentación
Fecha: Agosto 2025
"""

import os
import sys
from pathlib import Path
import markdown2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from weasyprint import HTML, CSS
import re

class ManualConverter:
    def __init__(self, base_path=None):
        self.base_path = base_path or Path.cwd()
        self.output_dir = self.base_path / "manuales_exportados"
        self.output_dir.mkdir(exist_ok=True)

    def markdown_to_word(self, markdown_file, output_file=None):
        """Convierte un archivo Markdown a Word (.docx)"""
        try:
            print(f"🔄 Convirtiendo {markdown_file} a Word...")

            # Leer el archivo markdown
            with open(markdown_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            # Crear documento Word
            doc = Document()

            # Configurar estilos
            self._setup_word_styles(doc)

            # Procesar contenido línea por línea
            lines = markdown_content.split('\n')
            current_paragraph = None

            for line in lines:
                line = line.strip()
                if not line:
                    if current_paragraph:
                        current_paragraph = None
                    continue

                # Títulos principales (# ##)
                if line.startswith('# '):
                    title = line.replace('# ', '')
                    p = doc.add_heading(title, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif line.startswith('## '):
                    title = line.replace('## ', '')
                    doc.add_heading(title, level=2)

                elif line.startswith('### '):
                    title = line.replace('### ', '')
                    doc.add_heading(title, level=3)

                elif line.startswith('#### '):
                    title = line.replace('#### ', '')
                    doc.add_heading(title, level=4)

                # Separadores
                elif line.startswith('---'):
                    doc.add_page_break()

                # Imágenes
                elif line.startswith('!['):
                    self._add_image_to_word(doc, line)

                # Listas con viñetas
                elif line.startswith('- '):
                    text = line.replace('- ', '')
                    text = self._clean_markdown_formatting(text)
                    p = doc.add_paragraph(text, style='List Bullet')

                # Listas numeradas
                elif re.match(r'^\d+\.', line):
                    text = re.sub(r'^\d+\.\s*', '', line)
                    text = self._clean_markdown_formatting(text)
                    p = doc.add_paragraph(text, style='List Number')

                # Texto normal
                elif line and not line.startswith('#'):
                    text = self._clean_markdown_formatting(line)
                    if text:
                        doc.add_paragraph(text)

            # Guardar archivo
            if not output_file:
                filename = Path(markdown_file).stem + '.docx'
                output_file = self.output_dir / filename

            doc.save(output_file)
            print(f"✅ Archivo Word creado: {output_file}")
            return str(output_file)

        except Exception as e:
            print(f"❌ Error al convertir a Word: {e}")
            return None

    def markdown_to_pdf(self, markdown_file, output_file=None):
        """Convierte un archivo Markdown a PDF"""
        try:
            print(f"🔄 Convirtiendo {markdown_file} a PDF...")

            # Leer el archivo markdown
            with open(markdown_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            # Procesar imágenes para rutas relativas
            markdown_content = self._fix_image_paths(markdown_content)

            # Convertir Markdown a HTML
            html_content = markdown2.markdown(
                markdown_content,
                extras=['fenced-code-blocks', 'tables', 'strike', 'task_list']
            )

            # Crear HTML completo con estilos CSS
            full_html = self._create_styled_html(html_content)

            # Definir archivo de salida
            if not output_file:
                filename = Path(markdown_file).stem + '.pdf'
                output_file = self.output_dir / filename

            # Convertir a PDF
            HTML(string=full_html, base_url=str(self.base_path)).write_pdf(output_file)
            print(f"✅ Archivo PDF creado: {output_file}")
            return str(output_file)

        except Exception as e:
            print(f"❌ Error al convertir a PDF: {e}")
            return None

    def _setup_word_styles(self, doc):
        """Configura estilos personalizados para el documento Word"""
        styles = doc.styles

        # Estilo para código
        try:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New'
            code_style.font.size = Inches(0.1)
        except:
            pass  # El estilo ya existe

    def _clean_markdown_formatting(self, text):
        """Limpia formato Markdown básico del texto"""
        # Negrita
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Cursiva
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Código inline
        text = re.sub(r'`(.*?)`', r'\1', text)
        # Enlaces
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return text

    def _add_image_to_word(self, doc, line):
        """Añade una imagen al documento Word"""
        try:
            # Extraer ruta de la imagen
            match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)

                # Convertir ruta relativa a absoluta
                if not os.path.isabs(image_path):
                    image_path = self.base_path / image_path

                # Verificar si la imagen existe
                if os.path.exists(image_path):
                    doc.add_picture(str(image_path), width=Inches(6))
                    if alt_text:
                        p = doc.add_paragraph(f"Figura: {alt_text}")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph(f"[Imagen no encontrada: {image_path}]")
        except Exception as e:
            print(f"⚠️  Error al añadir imagen: {e}")

    def _fix_image_paths(self, content):
        """Corrige las rutas de las imágenes para el PDF"""
        def replace_image_path(match):
            alt_text = match.group(1)
            image_path = match.group(2)

            # Convertir espacios en %20 para URLs
            image_path = image_path.replace(' ', '%20')

            # Si es ruta relativa, convertir a absoluta
            if not image_path.startswith('http') and not os.path.isabs(image_path):
                abs_path = self.base_path / image_path
                if abs_path.exists():
                    image_path = abs_path.as_uri()

            return f'![{alt_text}]({image_path})'

        return re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', replace_image_path, content)

    def _create_styled_html(self, html_content):
        """Crea HTML completo con estilos CSS para el PDF"""
        css_styles = """
        <style>
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                margin: 40px;
                color: #333;
                font-size: 12px;
            }
            h1 {
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
                page-break-before: always;
                font-size: 24px;
            }
            h2 {
                color: #34495e;
                border-bottom: 2px solid #95a5a6;
                padding-bottom: 5px;
                margin-top: 30px;
                font-size: 20px;
            }
            h3 {
                color: #7f8c8d;
                margin-top: 25px;
                font-size: 16px;
            }
            h4 {
                color: #95a5a6;
                margin-top: 20px;
                font-size: 14px;
            }
            p {
                margin-bottom: 15px;
                text-align: justify;
            }
            ul, ol {
                margin-bottom: 15px;
            }
            li {
                margin-bottom: 5px;
            }
            strong {
                color: #2c3e50;
            }
            code {
                background-color: #f8f9fa;
                padding: 2px 4px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 11px;
            }
            blockquote {
                border-left: 4px solid #3498db;
                padding-left: 20px;
                margin-left: 0;
                font-style: italic;
                background-color: #f8f9fa;
                padding: 15px;
            }
            img {
                max-width: 100%;
                height: auto;
                display: block;
                margin: 20px auto;
                border: 1px solid #ddd;
                border-radius: 5px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin-bottom: 20px;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }
            th {
                background-color: #3498db;
                color: white;
            }
            .page-break {
                page-break-before: always;
            }
            @page {
                margin: 2cm;
                @bottom-right {
                    content: "Página " counter(page);
                }
            }
        </style>
        """

        return f"""
        <!DOCTYPE html>
        <html lang="es">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Manual de Usuario</title>
            {css_styles}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

    def convert_all_manuals(self):
        """Convierte todos los manuales encontrados"""
        manual_files = [
            'Manual_Usuario_DesUr.md',
            'Manual_Usuario_CMIN.md'
        ]

        results = {
            'word': [],
            'pdf': [],
            'errors': []
        }

        for manual_file in manual_files:
            file_path = self.base_path / manual_file
            if file_path.exists():
                print(f"\n📖 Procesando: {manual_file}")

                # Convertir a Word
                word_result = self.markdown_to_word(file_path)
                if word_result:
                    results['word'].append(word_result)
                else:
                    results['errors'].append(f"Error al convertir {manual_file} a Word")

                # Convertir a PDF
                pdf_result = self.markdown_to_pdf(file_path)
                if pdf_result:
                    results['pdf'].append(pdf_result)
                else:
                    results['errors'].append(f"Error al convertir {manual_file} a PDF")
            else:
                results['errors'].append(f"Archivo no encontrado: {manual_file}")

        return results

def main():
    """Función principal"""
    print("🚀 Conversor de Manuales de Usuario")
    print("="*50)

    # Crear conversor
    converter = ManualConverter()

    # Convertir todos los manuales
    results = converter.convert_all_manuals()

    # Mostrar resultados
    print("\n📊 RESUMEN DE CONVERSIÓN")
    print("="*50)

    if results['word']:
        print("\n✅ ARCHIVOS WORD CREADOS:")
        for file in results['word']:
            print(f"   📄 {file}")

    if results['pdf']:
        print("\n✅ ARCHIVOS PDF CREADOS:")
        for file in results['pdf']:
            print(f"   📄 {file}")

    if results['errors']:
        print("\n❌ ERRORES:")
        for error in results['errors']:
            print(f"   ⚠️  {error}")

    print(f"\n📁 Archivos guardados en: {converter.output_dir}")
    print("\n✨ ¡Conversión completada!")

if __name__ == "__main__":
    main()
